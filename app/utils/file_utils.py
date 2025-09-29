"""
Enhanced File Utils with Better PDF/Document Parsing
Fixes common parsing issues and adds robust error handling.
"""

import os
import re
import tempfile
from pathlib import Path
from typing import Optional, Union

from app.core.exceptions import FileParsingException
from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_text_from_file(file_path: Union[str, Path]) -> str:
    """
    Enhanced text extraction with better error handling and fallbacks.
    
    Args:
        file_path: Path to the file to extract text from
        
    Returns:
        Extracted text content
        
    Raises:
        FileParsingException: If extraction fails
    """
    if not file_path or not os.path.exists(str(file_path)):
        raise FileParsingException(f"File not found: {file_path}")
    
    file_path = Path(file_path)
    file_extension = file_path.suffix.lower()
    
    logger.info(f"Extracting text from {file_extension} file: {file_path.name}")
    
    try:
        if file_extension == '.pdf':
            return _extract_pdf_text_enhanced(str(file_path))
        elif file_extension in ['.docx', '.doc']:
            return _extract_docx_text_enhanced(str(file_path))
        elif file_extension == '.txt':
            return _extract_txt_text_enhanced(str(file_path))
        else:
            raise FileParsingException(f"Unsupported file type: {file_extension}")
            
    except FileParsingException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error extracting text from {file_path}: {str(e)}")
        raise FileParsingException(f"Failed to extract text: {str(e)}")


def _extract_pdf_text_enhanced(file_path: Union[str, Path]) -> str:
    """
    Enhanced PDF text extraction with multiple fallback methods.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
        
    Raises:
        FileParsingException: If all extraction methods fail
    """
    text_content = ""
    extraction_errors = []
    
    # Method 1: pypdf (primary method)
    try:
        from pypdf import PdfReader
        
        logger.debug("Attempting PDF extraction with pypdf")
        
        with open(str(file_path), 'rb') as file:
            pdf_reader = PdfReader(file)
            
            # Handle encrypted PDFs
            if pdf_reader.is_encrypted:
                logger.warning(f"PDF {file_path} is encrypted, attempting to decrypt")
                try:
                    pdf_reader.decrypt("")  # Try empty password
                except Exception as decrypt_error:
                    raise FileParsingException(f"Cannot decrypt password-protected PDF: {str(decrypt_error)}")
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                        logger.debug(f"Extracted {len(page_text)} chars from page {page_num + 1}")
                except Exception as page_error:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(page_error)}")
                    continue
        
        if text_content.strip():
            logger.info(f"pypdf extraction successful: {len(text_content)} characters")
            return _clean_extracted_text(text_content)
        else:
            extraction_errors.append("pypdf: No text content extracted")
            
    except ImportError:
        extraction_errors.append("pypdf: Package not installed")
    except Exception as e:
        extraction_errors.append(f"pypdf: {str(e)}")
    
    # Method 2: Basic binary reading fallback
    try:
        logger.debug("Attempting basic binary reading fallback")
        
        with open(str(file_path), 'rb') as f:
            raw_content = f.read()
            
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
                try:
                    text_content = raw_content.decode(encoding, errors='ignore')
                    # Filter out binary junk, keep only printable text
                    text_content = ''.join(char for char in text_content if char.isprintable() or char.isspace())
                    
                    if len(text_content.strip()) > 50:  # Reasonable amount of text
                        logger.info(f"Binary fallback extraction successful with {encoding}: {len(text_content)} characters")
                        return _clean_extracted_text(text_content)
                        
                except UnicodeDecodeError:
                    continue
        
        extraction_errors.append("Binary fallback: No readable text found")
        
    except Exception as e:
        extraction_errors.append(f"Binary fallback: {str(e)}")
    
    # All methods failed
    error_summary = "; ".join(extraction_errors)
    raise FileParsingException(f"All PDF extraction methods failed: {error_summary}")


def _extract_docx_text_enhanced(file_path: Union[str, Path]) -> str:
    """
    Enhanced DOCX text extraction with error handling.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        FileParsingException: If extraction fails
    """
    try:
        from docx import Document
        
        logger.debug("Extracting text from DOCX file")
        
        doc = Document(str(file_path))
        text_content = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " "
                text_content += "\n"
        
        if not text_content.strip():
            raise FileParsingException("DOCX file contains no readable text")
        
        logger.info(f"DOCX extraction successful: {len(text_content)} characters")
        return _clean_extracted_text(text_content)
        
    except ImportError:
        raise FileParsingException("python-docx package not installed")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise FileParsingException(f"Failed to extract DOCX text: {str(e)}")


def _extract_txt_text_enhanced(file_path: Union[str, Path]) -> str:
    """
    Enhanced text file extraction with encoding detection.
    
    Args:
        file_path: Path to text file
        
    Returns:
        Extracted text content
        
    Raises:
        FileParsingException: If extraction fails
    """
    encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings_to_try:
        try:
            logger.debug(f"Trying to read text file with {encoding} encoding")
            
            with open(str(file_path), 'r', encoding=encoding) as file:
                text_content = file.read()
                
            if text_content.strip():
                logger.info(f"Text file extraction successful with {encoding}: {len(text_content)} characters")
                return _clean_extracted_text(text_content)
                
        except UnicodeDecodeError:
            logger.debug(f"Failed to decode with {encoding}, trying next encoding")
            continue
        except Exception as e:
            logger.error(f"Error reading text file with {encoding}: {str(e)}")
            continue
    
    raise FileParsingException(f"Failed to read text file with any encoding: {encodings_to_try}")


def _clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Normalize line breaks
    text = re.sub(r'\n+', '\n', text)
    
    # Remove non-printable characters except common whitespace
    text = ''.join(char for char in text if char.isprintable() or char in '\n\t ')
    
    # Strip and ensure we have content
    text = text.strip()
    
    # Validate minimum content
    if len(text) < 10:
        raise FileParsingException(f"Extracted text too short ({len(text)} chars): '{text[:50]}...'")
    
    logger.debug(f"Text cleaned: {len(text)} characters")
    return text


def save_uploaded_file(file_content: bytes, filename: str, temp_dir: str = "./data/temp") -> str:
    """
    Save uploaded file content to temporary directory with validation.
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        temp_dir: Temporary directory path
        
    Returns:
        Path to saved temporary file
        
    Raises:
        FileParsingException: If saving fails
    """
    try:
        # Create temp directory if it doesn't exist
        temp_path = Path(temp_dir)
        temp_path.mkdir(parents=True, exist_ok=True)
        
        # Validate file size (10MB limit)
        size_mb = len(file_content) / (1024 * 1024)
        if size_mb > 10:
            raise FileParsingException(f"File too large: {size_mb:.1f}MB (max 10MB)")
        
        # Validate filename
        if not filename or '..' in filename or filename.startswith('/'):
            raise FileParsingException(f"Invalid filename: {filename}")
        
        # Create unique temporary file
        safe_filename = "".join(c for c in filename if c.isalnum() or c in '._-')
        temp_file_path = temp_path / f"temp_{safe_filename}"
        
        # Ensure unique filename
        counter = 1
        while temp_file_path.exists():
            name_parts = safe_filename.rsplit('.', 1)
            if len(name_parts) == 2:
                base_name, extension = name_parts
                temp_file_path = temp_path / f"temp_{base_name}_{counter}.{extension}"
            else:
                temp_file_path = temp_path / f"temp_{safe_filename}_{counter}"
            counter += 1
        
        # Save file content
        temp_file_path.write_bytes(file_content)
        
        logger.info(f"Saved uploaded file to: {temp_file_path}")
        return str(temp_file_path)
        
    except Exception as e:
        logger.error(f"Failed to save uploaded file: {str(e)}")
        raise FileParsingException(f"Failed to save uploaded file: {str(e)}")


def cleanup_temp_file(file_path: Union[str, Path]) -> None:
    """
    Remove temporary file safely.
    
    Args:
        file_path: Path to temporary file
    """
    try:
        if file_path and os.path.exists(str(file_path)):
            os.unlink(str(file_path))
            logger.debug(f"Cleaned up temporary file: {file_path}")
    except Exception as e:
        logger.warning(f"Could not clean up temporary file {file_path}: {str(e)}")


def validate_file_type(filename: str, allowed_types: list) -> bool:
    """
    Validate if file type is allowed.
    
    Args:
        filename: Name of the file
        allowed_types: List of allowed file extensions (without dots)
        
    Returns:
        True if file type is allowed
    """
    if not filename:
        return False
    
    file_extension = Path(filename).suffix.lower().lstrip('.')
    return file_extension in [ext.lower() for ext in allowed_types]


def validate_file_size(file_content: bytes, max_size_mb: int) -> bool:
    """
    Validate file size.
    
    Args:
        file_content: File content in bytes
        max_size_mb: Maximum size in MB
        
    Returns:
        True if file size is within limit
    """
    size_mb = len(file_content) / (1024 * 1024)
    return size_mb <= max_size_mb